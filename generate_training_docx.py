"""Generate Yatharthata LIMS User Training Guide as a proper Word (.docx) document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style_title = doc.styles['Heading 1']
style_h2 = doc.styles['Heading 2']
style_h3 = doc.styles['Heading 3']
style_normal = doc.styles['Normal']

def add_heading(text, level=1):
    if level == 1:
        p = doc.add_heading(text, level=0)
    elif level == 2:
        p = doc.add_heading(text, level=1)
    else:
        p = doc.add_heading(text, level=2)
    return p

def add_para(text):
    return doc.add_paragraph(text)

def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_numbered(items):
    for i, item in enumerate(items, 1):
        doc.add_paragraph(item, style='List Number')

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = str(h)
        for p in hrow.cells[i].paragraphs:
            p.runs[0].bold = True
    for ri, row in enumerate(rows):
        r = t.rows[ri + 1]
        for ci, cell in enumerate(row):
            if ci < len(r.cells):
                r.cells[ci].text = str(cell)

# --- CONTENT ---
add_heading('Yatharthata LIMS Phase II – User Training Guide', 1)
add_para('Hydraulic Torque Wrench (HTW) Calibration Module')
add_para('Version: 1.0 | Last Updated: March 2026 | Applicable To: Admin, Engineer, and Customer roles')
doc.add_paragraph()

add_heading('1. Introduction', 2)
add_heading('1.1 Purpose', 3)
add_para('This guide explains how to use Yatharthata LIMS Phase II for Hydraulic Torque Wrench (HTW) calibration workflows. The system replaces manual Excel-based processes with a structured, role-based, and automated workflow aligned with ISO 6789-1 compliance and traceability.')

add_heading('1.2 Project Overview', 3)
add_para('Yatharthata LIMS Phase II automates the calibration workflow for Hydraulic Torque Wrenches (HTW). The system digitizes the complete lifecycle from job creation to certificate issuance, ensuring compliance with ISO 6789-1 standards while maintaining strict data control and operational traceability.')

add_heading('1.3 Key Capabilities', 3)
add_bullets([
    'Job creation based on Updated Inward',
    'Automated measurement data handling',
    'Master standard and nomenclature management',
    'Environmental validation',
    'Calibration calculations (automated)',
    'Uncertainty budget automation',
    'Certificate generation (CRT-1, CRT-2, CRT-3)',
    'License module (subscription-based access control)',
    'Usage report module (weekly and monthly reports)',
    'Role-based access control and audit trail',
])

add_heading('1.4 System Scope – Not Included', 3)
add_bullets(['Billing', 'Inventory / lifecycle management', 'Mobile application'])
doc.add_paragraph()

add_heading('2. System Access & Login', 2)
add_heading('2.1 Logging In', 3)
add_numbered([
    'Navigate to the application in your browser (e.g., http://localhost:3000).',
    'On the Login page, enter your email and password.',
    'Click Sign In.',
    'You will be redirected to your role-specific portal: Admin → Admin Portal (/admin); Engineer → Engineer Portal (/engineer); Customer → Customer Portal (/customer).',
])
add_heading('2.2 Forgot Password', 3)
add_para('Click Forgot password? on the login page to navigate to the password reset flow. Enter your email to receive reset instructions.')

add_heading('2.3 License Restrictions', 3)
add_para('If the system license is expired or inactive, a license modal will appear. License expired blocks login entirely—contact your system administrator. License status is validated at login.')
doc.add_paragraph()

add_heading('3. User Roles and Access', 2)
add_table(
    ['Role', 'Access Level', 'Key Responsibilities'],
    [
        ['Admin', 'Full system configuration', 'User management, master standards, uncertainty references, certificate approval, license oversight'],
        ['Engineer', 'Calibration operations', 'Create/update inwards, manage SRFs, perform calibrations, enter measurements, generate certificate drafts'],
        ['Customer', 'Read-only / status view', 'View SRFs, inwards, certificates; track status; review and respond to FIRs'],
    ]
)
add_heading('Engineer (Technician)', 3)
add_bullets(['Performs calibration', 'Enters measurement readings', 'Generates certificate drafts', 'Cannot modify formulas, master data, or standards'])
add_heading('Admin', 3)
add_bullets(['Manages master standards and reference data', 'Controls uncertainty references and CMC scope', 'Reviews and approves certificates', 'Manages license subscriptions', 'Oversees system configuration and access'])
doc.add_paragraph()

add_heading('4. Engineer Portal (Technicians)', 2)
add_heading('4.1 Engineer Dashboard', 3)
add_para('The Engineer Dashboard displays Quick Action cards, optional alerts for scheduled email reminders and failed notifications, and badge counts (e.g., drafts, reviewed FIRs) for items needing attention.')

add_heading('4.2 Quick Actions', 3)
add_para('Create Inward – Process incoming equipment and create SRF items. Use this to start a new inward from scratch.')
add_para('View & Update Inward – Manage existing inward entries and SRFs. Search, filter, edit, and update status. Badge indicates reviewed FIRs requiring action.')
add_para('Export Inward – Filter and export updated inward records.')
add_para('SRF Management – View and manage Service Request Forms. Navigate to SRF list and detail pages.')
add_para('Jobs Management – Manage calibration jobs and job status. Jobs are auto-created when an inward is updated/approved. Filter by Pending, In Progress, Completed, Terminated. Open a job and use "Start Calibration" to enter measurements.')
add_para('Calibration – Perform calibration data entry. Equipment and operating details auto-populate. Enter environmental conditions (temperature, humidity); validation blocks completion if criteria fail. Sections: Repeatability, Reproducibility, Output Drive, Drive Interface, Loading Point. Record locking prevents concurrent edits.')
add_para('Uncertainty Budget – View uncertainty budget for a job. Calculated by the system; read-only for Engineers.')
add_para('Certificates – Access certificate generation and management.')
add_para('View Deviations – Access deviation reports.')

add_heading('4.3 Scheduled Emails & Failed Notifications', 3)
add_para('Manage Scheduled Emails: view and send scheduled reminder emails. Review Failed Emails: retry failed notification deliveries.')
doc.add_paragraph()

add_heading('5. Admin Portal', 2)
add_heading('5.1 Admin Dashboard', 3)
add_para('Shows summary stats (total users, active, inactive) and quick actions: Invite New User, Manage Users, Master Standards.')

add_heading('5.2 Invite User', 3)
add_numbered([
    'Go to Invite User in the sidebar.',
    'Select role: Customer, Engineer, or Admin.',
    'For Customer: Enter company name, Ship To/Bill To addresses, contact name, phone, email.',
    'For Engineer/Admin: Enter full name and email.',
    'Click Send Invitation. The recipient receives an email to complete registration.',
])
add_para('For new companies: use Add New Company to create a company and associate the invitation.')

add_heading('5.3 User Management', 3)
add_para('View all users with filters (All, Admin, Engineer, Customer). Search by name, email, or company. Group customers by company. Activate/deactivate individual users. Activate/deactivate entire companies (batch).')

add_heading('5.4 Master Standards', 3)
add_para('Configure calibration reference data: Master Standard Details, Manufacturer Specifications, Pressure Gauge Resolution, Nomenclature & Range, Uncertainty References, CMC Reference, Coverage Factors, T-Distribution, Max Value Measurement Error, etc. Standards must be valid for the calibration range; expired standards block calibration.')

add_heading('5.5 Settings', 3)
add_para('Reserved for future global system settings.')
doc.add_paragraph()

add_heading('6. Customer Portal', 2)
add_heading('6.1 Customer Dashboard', 3)
add_para('Overview of Total SRFs, Active Deviations, Ready Certificates, Draft SRFs, FIRs for Review.')

add_heading('6.2 Main Functions', 3)
add_para('Track Status – Monitor equipment and SRF status.')
add_para('First Inspection Reports (FIRs) – List of inwards with completed first inspection. Review and provide feedback. Badge indicates items needing action.')
add_para('SRF View – View SRFs and their details.')
add_para('Customer Remarks Portal – Public access via token at /portal/inwards/:inwardId/remarks for adding remarks without full login.')
doc.add_paragraph()

add_heading('7. End-to-End Workflow Summary', 2)
add_numbered([
    'Inward Creation – Engineer creates inward and SRF items.',
    'Inward Update/Approval – Inward is updated and approved (Updated Inward).',
    'Job Creation – Jobs auto-created per equipment (1 job per equipment).',
    'Calibration – Engineer enters measurements and environmental data.',
    'Validation – System validates environmental conditions and standard validity.',
    'Calculation – System computes repeatability, reproducibility, uncertainty, etc.',
    'Certificate Draft – Engineer generates draft.',
    'Admin Review – Admin reviews and approves certificate.',
    'Issuance – Certificate issued to customer.',
])
doc.add_paragraph()

add_heading('8. Job States', 2)
add_bullets([
    'Pending – Job created, not yet started',
    'In Progress – Calibration in progress',
    'Completed – Calibration finished',
    'On Hold – Temporarily paused',
    'Terminated – Stopped/cancelled',
    'Not Calibrated – Equipment not calibrated',
])
doc.add_paragraph()

add_heading('9. Key Functional Modules (Technical Reference)', 2)
add_heading('9.1 Job Management', 3)
add_para('Job auto-created upon Updated Inward approval. One job per equipment. Job states: Pending, In Progress, Completed, On Hold, Terminated, Not Calibrated.')

add_heading('9.2 Measurement Module', 3)
add_para('Auto-populates equipment and operating details. Allows technicians to enter only required readings. Validates environmental conditions before processing. Blocks workflow if validation criteria are not met.')

add_heading('9.3 Master Standard Management', 3)
add_para('Automatic selection based on Model, Range, Nomenclature. Admin-controlled reference tables. Validity date enforcement. No calibration allowed with expired standards.')

add_heading('9.4 Calculation Engine', 3)
add_para('Automates repeatability, reproducibility, geometric effects, uncertainty evaluation. Logic aligned with validated Excel models. Results stored permanently for audit.')

add_heading('9.5 Certificate Module', 3)
add_para('Supports CRT-1, CRT-2, CRT-3. Draft creation by Engineer, review and approval by Admin, final issuance to customer. Includes device details, environmental conditions, calibration results, uncertainty values, ULR (editable). Full audit trail maintained.')

add_heading('9.6 License Module', 3)
add_para('Subscription-based access control. Company-level license validation. Expiry-based feature control. Usage tracking per organization.')

add_heading('9.7 Usage Report Module', 3)
add_para('Weekly and monthly usage reports. Job activity summaries. User activity metrics.')
doc.add_paragraph()

add_heading('10. Data & Compliance Controls', 2)
add_bullets([
    'Role-based access enforcement',
    'Company-level data isolation',
    'Record-level edit control',
    'Audit trail maintained for minimum 5 years',
    'Raw measurement data permanently stored',
    'All logic aligned with ISO compliance standards',
])
doc.add_paragraph()

add_heading('11. System Constraints', 2)
add_bullets([
    'Calculation logic must match validated Excel references',
    'Expired standards automatically block calibration',
    'Technicians cannot modify formulas or master data',
    'Certificate issuance requires completed job validation',
    'All workflow transitions are logged',
])
doc.add_paragraph()

add_heading('12. Training Session Plan (By Day)', 2)
add_heading('Day 1: Introduction, Access & Inward Basics', 3)
add_para('Topics: System overview; roles; login and navigation; Engineer Portal overview; Create Inward; View & Update Inward; Customer Portal; Customer Remarks Portal.')
add_para('Learning Outcomes: Users understand product purpose, roles, login flow; can create and update inwards; can use Customer Portal and Remarks.')
add_para('Activities: Hands-on login; create sample inward; edit inward; demo Customer Portal; use token URL for Remarks.')

add_heading('Day 2: SRF, Jobs & Calibration Workflow', 3)
add_para('Topics: SRF Management; Jobs Management (creation from Updated Inward, job states); Calibration Module; Environmental Validation; Uncertainty Budget.')
add_para('Learning Outcomes: Users understand SRF–Job–Calibration flow; can run calibration; know how validation works; can interpret Uncertainty Budget.')
add_para('Activities: SRF exploration; job listing and filtering; full calibration exercise; validation failure/correction; Uncertainty Budget view.')

add_heading('Day 3: Administration & Advanced Features', 3)
add_para('Topics: Admin Portal – User Management; Invite User; Master Standards; Uncertainty & CMC configuration; Certificate workflow (draft, review, approve); Export Inward; Reports and deviations; Constraints and best practices.')
add_para('Learning Outcomes: Admins can invite users and manage access; configure master standards; complete certificate workflow; use reports and export.')
add_para('Activities: Invite user; activate/deactivate; company grouping; add/update master standard; run certificate workflow; export and report demos.')

add_heading('Summary by Role', 3)
add_table(
    ['Day', 'Focus', 'Roles', 'Main Modules'],
    [
        ['1', 'Introduction & Inward', 'All', 'Login, Engineer Portal, Inward, Customer Portal, Remarks'],
        ['2', 'Calibration Workflow', 'Engineer (primary)', 'SRF, Jobs, Calibration, Validation, Uncertainty Budget'],
        ['3', 'Administration & Compliance', 'Admin (primary), Engineer', 'User Management, Master Standards, Certificates, Reports'],
    ]
)
doc.add_paragraph()

add_heading('13. Training Tips', 2)
add_bullets([
    'Role-based paths: Customers may attend Day 1 only; Engineers Days 1–2 plus selected Day 3; Admins full program.',
    'Environment: Use test accounts and sample inwards with approved equipment before Day 2.',
    'Order: Ensure inwards reach Updated Inward status before Day 2 so jobs exist.',
    'Handouts: Provide this guide plus role-specific quick-reference sheets.',
    'Recordings: Record Day 3 (Admin) for future reference.',
])
doc.add_paragraph()

add_heading('14. Troubleshooting', 2)
add_table(
    ['Issue', 'Action'],
    [
        ['Cannot log in', 'Verify credentials; check account is active; confirm license status'],
        ['License expired', 'Contact system administrator'],
        ['Environmental validation failed', 'Recheck temperature and humidity values'],
        ['Calibration blocked by standards', 'Verify master standard validity dates (Admin)'],
        ['Record locked', 'Another user may be editing; wait or retry'],
        ['Forgot password', 'Use Forgot password link on login page'],
    ]
)
doc.add_paragraph()

add_heading('15. Support Contacts', 2)
add_para('License or access issues: Contact AIMLSN YatharthataLIMS System Administrator.')
add_para('Technical support: Contact your organization\'s IT support or system administrator.')
doc.add_paragraph()
add_para('This training guide is based on Yatharthata LIMS Phase II version 1.0 and reflects the system as of March 2026.')

# Save
out_path = r'c:\Users\Pranav Kulkarni\Downloads\LIMSCPIIV - Copy 1\LIMSCPIIV - Copy\LIMSZ\lims-phase-2\Yatharthata_LIMS_User_Training_Guide.docx'
doc.save(out_path)
print(f'Created: {out_path}')
