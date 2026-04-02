import os 

from docx import Document 

 
 

doc = Document() 

folder = "."  # project root 

 
 

skip_dirs = {"node_modules", ".git", "dist", "build", "venv", "__pycache__"} 

 
 

for root, dirs, files in os.walk(folder): 

    # Sort dirs and files alphabetically 

    dirs[:] = sorted([d for d in dirs if d not in skip_dirs]) 

    files = sorted(files) 

 
 

    for file in files: 

        if file.endswith((".ts", ".tsx", ".js",".css",".html",".json")):  # frontend source files 

            filepath = os.path.join(root, file) 

            rel_path = os.path.relpath(filepath, folder) 

 
 

            # Add file path as heading 

            doc.add_heading(rel_path, level=2) 

 
 

            try: 

                with open(filepath, "r", encoding="utf-8", errors="ignore") as f: 

                    code = f.read() 

                doc.add_paragraph(code, style="Courier New") 

            except Exception as e: 

                doc.add_paragraph(f"<<Could not read file: {e}>>") 

 
 

doc.save("frontend.docx") 

 
 