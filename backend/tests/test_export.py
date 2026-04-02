import importlib
import sys
from pathlib import Path

import docx


class FakeDoc:
    def __init__(self) -> None:
        self.headings = []
        self.paragraphs = []
        self.saved_path = None

    def add_heading(self, text: str, level: int) -> None:
        self.headings.append((text, level))

    def add_paragraph(self, text: str, style: str | None = None) -> None:
        self.paragraphs.append((text, style))

    def save(self, path: str) -> None:
        self.saved_path = Path(path)


def test_export_documents_python_sources(tmp_path, monkeypatch):
    fake_doc = FakeDoc()

    # Ensure the export module walks the test workspace rather than the real project
    monkeypatch.chdir(tmp_path)

    # Provide a predictable Document factory so we can inspect the interactions
    monkeypatch.setattr(docx, "Document", lambda: fake_doc)

    # Create sample files that should and should not be included in the export
    (tmp_path / "root.py").write_text('print("root")', encoding="utf-8")

    src_dir = tmp_path / "src"
    src_dir.mkdir()
    (src_dir / "module.py").write_text('print("module")', encoding="utf-8")

    skip_dir = tmp_path / "venv"
    skip_dir.mkdir()
    (skip_dir / "ignored.py").write_text("ignored = True", encoding="utf-8")

    # Reload the export module so the top-level script executes inside the sandbox
    sys.modules.pop("backend.export", None)
    importlib.import_module("backend.export")

    # The script should write to the expected document name
    assert fake_doc.saved_path == Path("backend.docx")

    # It should capture headings for the Python files we created (but not the skipped ones)
    expected_headings = [
        ("root.py", 2),
        (str(Path("src") / "module.py"), 2),
    ]
    assert fake_doc.headings == expected_headings

    # And store the corresponding source code in the paragraphs
    paragraphs_text = [text for text, _ in fake_doc.paragraphs]
    assert 'print("root")' in paragraphs_text
    assert 'print("module")' in paragraphs_text

    # Verify that files from skipped directories were not processed
    for heading, _ in fake_doc.headings:
        assert not heading.startswith("venv")

