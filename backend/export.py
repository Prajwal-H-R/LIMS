import os 

from docx import Document 

 
 

doc = Document() 

folder = "."  # repo root 

 
 

skip_dirs = {"venv", ".git", "node_modules", "__pycache__", "dist", "build"} 

 
 

for root, dirs, files in os.walk(folder): 

    # Sort dirs and files alphabetically 

    dirs[:] = sorted([d for d in dirs if d not in skip_dirs]) 

    files = sorted(files) 

 
 

    for file in files: 

        if file.endswith(".py"):  # only Python files 

            filepath = os.path.join(root, file) 

            rel_path = os.path.relpath(filepath, folder) 

 
 

            # Add file path as heading (shows folders too) 

            doc.add_heading(rel_path, level=2) 

 
 

            try: 

                with open(filepath, "r", encoding="utf-8", errors="ignore") as f: 

                    code = f.read() 

                doc.add_paragraph(code, style="Courier New") 

            except Exception as e: 

                doc.add_paragraph(f"<<Could not read file: {e}>>") 

 
 

doc.save("backend.docx") 

 
 